import pdfplumber
from openpyxl import load_workbook
import docx
from typing import Dict, Any
import os

class DocumentProcessor:
    @staticmethod
    def process(file_path: str) -> Dict[str, Any]:
        """
        Procesa un documento y extrae su contenido y metadata.
        
        Args:
            file_path: Ruta al archivo a procesar
            
        Returns:
            Dict con el texto extraído y metadata del documento
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"El archivo {file_path} no existe")
            
        file_type = file_path.split('.')[-1].lower()
        
        try:
            if file_type == 'pdf':
                return DocumentProcessor._process_pdf(file_path)
            elif file_type in ['xlsx', 'xls']:
                return DocumentProcessor._process_excel(file_path)
            elif file_type in ['docx', 'doc']:
                return DocumentProcessor._process_word(file_path)
            else:
                raise ValueError(f"Formato de archivo no soportado: {file_type}")
        except Exception as e:
            raise Exception(f"Error procesando {file_type}: {str(e)}")

    @staticmethod
    def _process_pdf(file_path: str) -> Dict[str, Any]:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""
            
            return {
                "texto": text,
                "metadata": {
                    "paginas": len(pdf.pages),
                    "tamano": os.path.getsize(file_path)
                }
            }

    @staticmethod
    def _process_excel(file_path: str) -> Dict[str, Any]:
        wb = load_workbook(file_path, data_only=True)
        sheets_data = {}
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_data = []
            for row in ws.rows:
                row_data = [str(cell.value) if cell.value is not None else "" for cell in row]
                sheet_data.append(row_data)
            sheets_data[sheet_name] = sheet_data

        return {
            "texto": "",  # Excel no tiene texto continuo
            "metadata": {
                "hojas": sheets_data,
                "num_hojas": len(wb.sheetnames),
                "tamano": os.path.getsize(file_path)
            }
        }

    @staticmethod
    def _process_word(file_path: str) -> Dict[str, Any]:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        return {
            "texto": text,
            "metadata": {
                "parrafos": len(doc.paragraphs),
                "tamano": os.path.getsize(file_path)
            }
        }
